"""Script para guardar una conversación en un documento Word."""

import os
from docx import Document

def guardar_chat_en_docx(texto_chat: str,
                          plantilla: str = "plantilla_base.docx",
                          salida: str = "chat_guardado.docx") -> None:
    """Carga la plantilla, procesa el chat y guarda el documento.

    Parameters
    ----------
    texto_chat : str
        Texto de la conversación a agregar.
    plantilla : str, optional
        Nombre de la plantilla a cargar. Por defecto ``plantilla_base.docx``.
    salida : str, optional
        Nombre del documento resultante. Por defecto ``chat_guardado.docx``.
    """
    if not os.path.exists(plantilla):
        raise FileNotFoundError(f"No se encontró la plantilla: {plantilla}")

    # Cargar documento base
    doc = Document(plantilla)

    # Procesar cada línea de la conversación
    for linea in texto_chat.splitlines():
        linea = linea.strip()
        if not linea:
            continue

        if linea.startswith("Usuario:"):
            hablante = "Usuario"
            mensaje = linea[len("Usuario:"):].strip()
        elif linea.startswith("ChatGPT:"):
            hablante = "ChatGPT"
            mensaje = linea[len("ChatGPT:"):].strip()
        else:
            # Ignorar líneas que no estén etiquetadas correctamente
            continue

        # Nombre del hablante en negrita
        p_nombre = doc.add_paragraph()
        run = p_nombre.add_run(f"{hablante}:")
        run.bold = True

        # Mensaje del hablante en un párrafo independiente
        doc.add_paragraph(mensaje)

    # Guardar el documento resultante
    doc.save(salida)


if __name__ == "__main__":
    # Texto de ejemplo simulando una conversación
    texto_chat = (
        "Usuario: Hola, ¿cómo estás?\n"
        "ChatGPT: ¡Hola! Estoy aquí para ayudarte.\n"
        "Usuario: Gracias por la información."
    )

    guardar_chat_en_docx(texto_chat)
